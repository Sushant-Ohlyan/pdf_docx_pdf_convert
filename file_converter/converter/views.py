from django.shortcuts import render

# Create your views here.
from django.shortcuts import render
from django.http import FileResponse, HttpResponse
from django.core.files.storage import FileSystemStorage
from docx import Document
from reportlab.pdfgen import canvas
import io
import os
import img2pdf
from PIL import Image
from pdf2docx import Converter
import fitz 

def home(request):
    return render(request, 'converter/home.html')


def docx_to_pdf(request):
    if request.method == 'POST':
        file = request.FILES['docx_file']
        doc = Document(file)
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer)
        for para in doc.paragraphs:
            pdf.drawString(100, 800 - doc.paragraphs.index(para)*20, para.text)
        pdf.save()
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename='converted.pdf')
    return render(request, 'converter/home.html')


def image_to_pdf(request):
    if request.method == 'POST':
        image = request.FILES['image_file']
        image = Image.open(image)
        pdf_bytes = img2pdf.convert(image.filename)
        return FileResponse(io.BytesIO(pdf_bytes), as_attachment=True, filename='output.pdf')
    return render(request, 'converter/home.html')



def pdf_to_docx(request):
    if request.method == 'POST':
        pdf_file = request.FILES['pdf_file']
        temp_path = f'media/{pdf_file.name}'
        with open(temp_path, 'wb+') as dest:
            for chunk in pdf_file.chunks():
                dest.write(chunk)
        docx_path = temp_path.replace('.pdf', '.docx')
        cv = Converter(temp_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return FileResponse(open(docx_path, 'rb'), as_attachment=True, filename='converted.docx')

 

def pdf_to_text(request):
    if request.method == 'POST':
        pdf_file = request.FILES['pdf_file']
        doc = fitz.open(stream=pdf_file.read(), filetype='pdf')
        text = "\n".join(page.get_text() for page in doc)
        response = HttpResponse(text, content_type='text/plain')
        response['Content-Disposition'] = 'attachment; filename="output.txt"'
        return response
    return render(request, 'converter/home.html')
